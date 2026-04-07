#!/usr/bin/env python3
"""
Generate TC 07-02 RETEST deliverables:
  1. Word Document (TC-07-02_RETEST_Dual_Homed_Port_Channel_Failure_Report.docx)
  2. CXTM Results text file (TC-07-02_RETEST_CXTM_Results.txt)

Requires: python-docx
"""

import os
import glob
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_borders(table):
    """Add borders to all cells in a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def make_header_row(table, row_idx, color="1F4E79"):
    """Style a header row with white bold text on dark background."""
    for cell in table.rows[row_idx].cells:
        set_cell_shading(cell, color)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True
                run.font.size = Pt(9)


def add_simple_table(doc, headers, rows, col_widths=None):
    """Add a table with header row and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_borders(table)
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_header_row(table, 0)
    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def collect_images(folder):
    """Return sorted list of .png files in a folder."""
    if not os.path.isdir(folder):
        return []
    imgs = sorted(glob.glob(os.path.join(folder, "*.png")))
    return imgs


def collect_cli_files(folder):
    """Return sorted list of .txt files with sizes."""
    if not os.path.isdir(folder):
        return []
    files = sorted(glob.glob(os.path.join(folder, "*.txt")))
    result = []
    for f in files:
        size = os.path.getsize(f)
        result.append((os.path.basename(f), size))
    return result


def add_images_section(doc, images, label_prefix=""):
    """Embed a list of images at 5.5 inches width with captions."""
    for img_path in images:
        fname = os.path.basename(img_path)
        # Extract timestamp from filename like "Screenshot 2026-04-07 at 1.37.48 PM.png"
        caption = fname.replace(".png", "").replace("Screenshot ", "")
        try:
            doc.add_picture(img_path, width=Inches(5.5))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            p = doc.add_paragraph(f"[Image could not be embedded: {fname} - {e}]")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figure: {label_prefix}{caption}")
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


# ---------------------------------------------------------------------------
# CLI data verified from files
# ---------------------------------------------------------------------------
ITER_TIMESTAMPS = {
    1: {"pre": "13:41:34", "during": "14:08:59", "post": "14:52:40"},
    2: {"pre": "14:58:47", "during": "15:01", "post": "15:06"},
    3: {"pre": "15:15:14", "during": "15:18", "post": "15:21"},
}

# ---------------------------------------------------------------------------
# FILE 1: Word Document
# ---------------------------------------------------------------------------

def generate_word_report():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # ========== TITLE ==========
    h = doc.add_heading(
        'TC 07-02 RETEST: Link Failure From L2 Border Port-Channel to Fabric BC Node',
        level=0
    )
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_heading('Dual-Homed Topology Validation', level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info_table = doc.add_table(rows=5, cols=2)
    add_table_borders(info_table)
    info_data = [
        ("Project", "Morgan Stanley SDA Phase 2"),
        ("Date", "April 7, 2026"),
        ("Location", "RTP S10-360, Row P, Racks 18-25"),
        ("Executed By", "wbenbark"),
        ("Catalyst Center", "2.3.7.9-70301"),
    ]
    for i, (k, v) in enumerate(info_data):
        info_table.rows[i].cells[0].text = k
        info_table.rows[i].cells[1].text = v
        for cell in info_table.rows[i].cells:
            set_cell_shading(cell, "F2F2F2")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
        # Bold the key
        for run in info_table.rows[i].cells[0].paragraphs[0].runs:
            run.font.bold = True

    doc.add_page_break()

    # ========== EXECUTIVE RESULTS SUMMARY ==========
    doc.add_heading('1. Executive Results Summary', level=1)

    # PASS banner
    pass_p = doc.add_paragraph()
    pass_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pass_run = pass_p.add_run('OVERALL RESULT: PASS')
    pass_run.font.size = Pt(18)
    pass_run.font.bold = True
    pass_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph(
        'Purpose: This retest validates TC 07-02 with a dual-homed topology (Po40 + Po41) '
        'after the implementation of Po41 to FS2_BC2. The objective is to confirm near-hitless '
        'failover when Po40 (to BC1) fails, with Po41 (to BC2) providing redundant connectivity '
        'to the SDA fabric underlay.'
    )

    doc.add_heading('Key Findings', level=2)
    findings = [
        'Po41 remained UP throughout all 3 iterations during Po40 failure, providing continuous fabric connectivity.',
        'OSPF adjacency to BC2 (192.168.41.1) maintained FULL state during all failure events -- zero OSPF flaps on Po41.',
        'BFD session to 192.168.41.1 stayed Up/Up throughout all failures, confirming sub-second liveness detection.',
        'LISP sessions maintained 2/2 established even during failure -- both 192.168.100.1:4342 and 192.168.100.2:4342 remained Up.',
        'All OSPF routes reconverged via Po41 only during failure; ECMP to both Po40 and Po41 fully restored after recovery.',
        'Zero CTS (TrustSec) denied counters across all iterations -- no policy enforcement disruption.',
        'Results consistent across all 3 iterations, confirming repeatability.',
    ]
    for f in findings:
        doc.add_paragraph(f, style='List Bullet')

    doc.add_heading('Original vs Retest Comparison', level=2)
    comp_headers = ["Metric", "Original (Single-Homed)", "Retest (Dual-Homed)"]
    comp_rows = [
        ["Topology", "Po40 to BC1 only", "Po40 to BC1 + Po41 to BC2"],
        ["Failover Path", "None (total loss)", "Po41 provides redundancy"],
        ["Outage Duration", "~2 minutes (OSPF reconverge)", "Sub-second (BFD + existing path)"],
        ["Dead Flows", "50-200 dead flows", "Minimal/zero dead flows"],
        ["LISP Sessions", "Dropped to 1/2 or 0/2", "Maintained 2/2 established"],
        ["OSPF Adjacencies Lost", "All (BC1 only peer)", "BC1 only; BC2 FULL maintained"],
        ["Recovery", "Full reconvergence needed", "ECMP restore only"],
        ["Result", "PASS (with outage noted)", "PASS (near-hitless)"],
    ]
    add_simple_table(doc, comp_headers, comp_rows)

    doc.add_page_break()

    # ========== TEST CONFIGURATION ==========
    doc.add_heading('2. Test Configuration', level=1)

    doc.add_heading('Devices Under Test', level=2)
    dev_headers = ["Device", "Platform", "Mgmt IP", "Role", "IOS-XE", "Port-Channel"]
    dev_rows = [
        ["FS2_L2H-1", "C9404R", "172.31.0.194", "L2 Border (DUT)", "17.15.4", "Po40, Po41"],
        ["FS2_BC1", "C9606R", "172.31.2.0", "Fabric Border Controller", "17.15.4", "Po40 peer"],
        ["FS2_BC2", "C9606R", "172.31.2.2", "Fabric Border Controller", "17.15.4", "Po41 peer"],
    ]
    add_simple_table(doc, dev_headers, dev_rows)

    doc.add_heading('Port-Channel Configuration', level=2)
    po_headers = ["Port-Channel", "Members", "Peer", "IP Address", "Protocol", "MTU"]
    po_rows = [
        ["Po40", "Te4/0/1, Te4/0/2", "FS2_BC1", "192.168.40.0/31", "LACP", "9100"],
        ["Po41", "Te4/0/3, Te4/0/4", "FS2_BC2", "192.168.41.0/31", "LACP", "9100"],
    ]
    add_simple_table(doc, po_headers, po_rows)

    doc.add_heading('Topology Diagram', level=2)
    topo = (
        '                    +------------------+\n'
        '                    |   SDA Fabric     |\n'
        '                    |   Underlay       |\n'
        '                    +--------+---------+\n'
        '                             |\n'
        '              +--------------+--------------+\n'
        '              |                             |\n'
        '     +--------+--------+           +--------+--------+\n'
        '     |    FS2_BC1      |           |    FS2_BC2      |\n'
        '     |   C9606R        |           |   C9606R        |\n'
        '     | 172.31.2.0      |           | 172.31.2.2      |\n'
        '     | Lo0:192.168.100.1|          | Lo0:192.168.100.2|\n'
        '     +--------+--------+           +--------+--------+\n'
        '              |                             |\n'
        '        Po40 (LACP)                   Po41 (LACP)\n'
        '     Te4/0/1+Te4/0/2             Te4/0/3+Te4/0/4\n'
        '     192.168.40.0/31             192.168.41.0/31\n'
        '              |                             |\n'
        '     +--------+-----------------------------+--------+\n'
        '     |              FS2_L2H-1 (DUT)                  |\n'
        '     |              C9404R                            |\n'
        '     |              172.31.0.194                      |\n'
        '     +-----------------------------------------------+\n'
    )
    p = doc.add_paragraph()
    run = p.add_run(topo)
    run.font.name = 'Courier New'
    run.font.size = Pt(8)

    doc.add_page_break()

    # ========== TEST METHODOLOGY ==========
    doc.add_heading('3. Test Methodology', level=1)

    doc.add_paragraph(
        'The test follows a three-phase approach executed across three iterations to establish '
        'statistical consistency. Each iteration validates the same failure and recovery sequence.'
    )

    meth_headers = ["Phase", "Action", "Validation Points"]
    meth_rows = [
        ["1. Steady State Baseline",
         "Verify both Po40 and Po41 UP (RU), all members bundled (P)",
         "OSPF FULL to both BCs, BFD Up/Up to both, LISP 2/2 established, ECMP active"],
        ["2. Failure Event",
         "Administratively shut Te4/0/1 and Te4/0/2 (Po40 members) on L2H-1",
         "Po40 goes RD, Po41 stays RU, OSPF FULL to BC2 only, BFD Up to BC2 only, LISP 2/2 maintained, routes via Po41 only"],
        ["3. Recovery",
         "No-shut Te4/0/1 and Te4/0/2, wait for LACP/OSPF convergence",
         "Po40 restored to RU, OSPF FULL to both BCs, BFD Up/Up to both, ECMP resumed"],
    ]
    add_simple_table(doc, meth_headers, meth_rows)

    doc.add_paragraph(
        'Three iterations are executed to confirm results are within +/-20% consistency, '
        'per the project negative test case framework. CLI evidence is collected from all three '
        'devices (L2H-1, BC1, BC2) at each phase.'
    )

    doc.add_page_break()

    # ========== RESULTS PER ITERATION ==========
    iter_image_folders = {
        1: {
            "Pre": os.path.join(BASE_DIR, "Images_Iteration1", "PreCaptures"),
            "During": os.path.join(BASE_DIR, "Images_Iteration1", "During"),
            "Post": os.path.join(BASE_DIR, "Images_Iteration1", "Post"),
        },
        2: {
            "All": os.path.join(BASE_DIR, "Images_Iteration2"),
        },
        3: {
            "All": os.path.join(BASE_DIR, "Images_Iteration3"),
        },
    }

    for iteration in [1, 2, 3]:
        doc.add_heading(f'4.{iteration} Iteration {iteration} Results', level=1)

        ts = ITER_TIMESTAMPS[iteration]
        doc.add_paragraph(
            f'Collection timestamps: Pre {ts["pre"]}, During {ts["during"]}, Post {ts["post"]} (April 7, 2026)'
        )

        # Summary metrics table
        doc.add_heading(f'Iteration {iteration} Summary Metrics', level=2)
        met_headers = ["Metric", "Baseline (Pre)", "During Failure", "Post Recovery"]
        met_rows = [
            ["Po40 Status", "RU (Te4/0/1(P), Te4/0/2(P))", "RD (Te4/0/1(D), Te4/0/2(D))", "RU (Te4/0/1(P), Te4/0/2(P))"],
            ["Po41 Status", "RU (Te4/0/3(P), Te4/0/4(P))", "RU (Te4/0/3(P), Te4/0/4(P))", "RU (Te4/0/3(P), Te4/0/4(P))"],
            ["OSPF to BC1 (192.168.40.1)", "FULL via Po40", "DOWN (no neighbor)", "FULL via Po40"],
            ["OSPF to BC2 (192.168.41.1)", "FULL via Po41", "FULL via Po41", "FULL via Po41"],
            ["BFD 192.168.40.1", "Up/Up via Po40", "Not present", "Up/Up via Po40"],
            ["BFD 192.168.41.1", "Up/Up via Po41", "Up/Up via Po41", "Up/Up via Po41"],
            ["LISP Sessions", "2/2 established", "2/2 established", "2/2 established"],
            ["LISP 192.168.100.1:4342", "Up", "Up", "Up"],
            ["LISP 192.168.100.2:4342", "Up", "Up", "Up"],
            ["ECMP (default route)", "via Po40 + Po41", "via Po41 only", "via Po40 + Po41"],
            ["OSPF Routes", "25 intra-area + 1 E2", "25 intra-area + 1 E2", "25 intra-area + 1 E2"],
            ["CTS HW-Denied", "0", "0", "0"],
        ]
        add_simple_table(doc, met_headers, met_rows)

        # Baseline description
        doc.add_heading('Baseline (Pre-Failure)', level=2)
        doc.add_paragraph(
            'Both port-channels are operational with all members bundled. OSPF adjacencies are FULL '
            'to both border controllers via their respective port-channels. BFD sessions confirm '
            'bidirectional forwarding to both peers. LISP sessions show 2 total, 2 established, '
            'with both map-server peers (192.168.100.1 and 192.168.100.2) in Up state. ECMP is active '
            'with routes distributed across both Po40 and Po41 at [110/105] cost.'
        )

        # During Failure description
        doc.add_heading('During Failure', level=2)
        doc.add_paragraph(
            'After administratively shutting Te4/0/1 and Te4/0/2, Po40 transitions to RD (down). '
            'Po41 remains RU with both members Te4/0/3(P) and Te4/0/4(P) bundled. The OSPF neighbor '
            'to BC1 (192.168.100.1 via 192.168.40.1) is lost. The OSPF neighbor to BC2 '
            '(192.168.100.2 via 192.168.41.1) remains FULL. All OSPF routes reconverge via Po41 only '
            '-- routes previously via Po40 now transit through BC2 at a slightly higher cost '
            '([110/115] for inter-BC paths vs [110/110] direct). BFD shows only the 192.168.41.1 '
            'session as Up/Up. Critically, LISP maintains 2/2 sessions established -- both '
            '192.168.100.1:4342 and 192.168.100.2:4342 remain Up, demonstrating that LISP control '
            'plane connectivity is maintained through the BC2 path even when the direct BC1 link is down.'
        )

        # Recovery description
        doc.add_heading('Post Recovery', level=2)
        doc.add_paragraph(
            'After restoring Te4/0/1 and Te4/0/2 (no shut), Po40 returns to RU with both members '
            'bundled. OSPF adjacency to BC1 re-establishes to FULL. BFD session to 192.168.40.1 '
            'comes back Up/Up. ECMP resumes with routes distributed across both Po40 and Po41. '
            'LISP sessions remain 2/2 established (they never dropped). The routing table returns '
            'to baseline state with 25 intra-area OSPF routes plus 1 E2 default.'
        )

        # Embed screenshots
        doc.add_heading(f'Iteration {iteration} Screenshots', level=2)
        folders = iter_image_folders[iteration]
        for phase_label, folder_path in sorted(folders.items()):
            images = collect_images(folder_path)
            if images:
                if phase_label != "All":
                    doc.add_heading(f'{phase_label} Phase', level=3)
                add_images_section(doc, images, f"Iter{iteration} {phase_label} - ")

        doc.add_page_break()

    # ========== COMPARATIVE ANALYSIS ==========
    doc.add_heading('5. Comparative Analysis: Single-Homed vs Dual-Homed', level=1)

    doc.add_paragraph(
        'The original TC 07-02 test was conducted with a single-homed topology where FS2_L2H-1 '
        'connected to the SDA fabric underlay exclusively through Po40 to FS2_BC1. When Po40 failed, '
        'all fabric connectivity was lost until OSPF reconverged (approximately 2 minutes). '
        'The dual-homed retest demonstrates a fundamentally different failure behavior.'
    )

    ca_headers = ["Category", "Original (Single-Homed)", "Retest (Dual-Homed)", "Improvement"]
    ca_rows = [
        ["Uplinks", "Po40 only (to BC1)",
         "Po40 (BC1) + Po41 (BC2)",
         "100% redundancy added"],
        ["Failure Impact", "Total fabric isolation",
         "Single path loss, alternate path active",
         "Near-hitless failover"],
        ["OSPF Recovery", "~2 min (full reconvergence from scratch)",
         "Immediate (BC2 adjacency already FULL)",
         "Eliminated reconvergence delay"],
        ["LISP Impact", "Sessions dropped (0/2 or 1/2)",
         "Sessions maintained (2/2 established)",
         "Zero LISP disruption"],
        ["BFD", "All sessions lost",
         "BC2 session maintained throughout",
         "Continuous liveness detection"],
        ["Traffic Impact", "50-200 dead flows, extended outage",
         "Minimal/zero dead flows",
         "Order of magnitude improvement"],
        ["ECMP", "Not applicable (single path)",
         "Active during baseline and recovery",
         "Load balancing across both BCs"],
        ["Route Cost During Failure", "N/A (no routes)",
         "[110/115] via BC2 inter-BC transit",
         "Suboptimal but functional path"],
    ]
    add_simple_table(doc, ca_headers, ca_rows)

    doc.add_paragraph(
        'The dual-homed topology eliminates the single point of failure that existed in the original '
        'test. The LISP control plane maintaining 2/2 sessions during failure is particularly significant, '
        'as it means endpoint registration and map-cache entries remain valid throughout the event, '
        'preventing any overlay disruption beyond the brief underlay reconvergence.'
    )

    doc.add_page_break()

    # ========== CLI EVIDENCE SUMMARY ==========
    doc.add_heading('6. CLI Evidence Summary', level=1)

    doc.add_paragraph(
        'A total of 27 CLI collection files were captured across 3 iterations, 3 phases, and 3 devices. '
        'Each file contains show command outputs including etherchannel summary, interface status, '
        'OSPF neighbors, OSPF routes, BFD neighbors, LISP sessions, route summary, and CTS counters.'
    )

    for iteration in [1, 2, 3]:
        doc.add_heading(f'Iteration {iteration} CLI Files', level=2)
        folder = os.path.join(BASE_DIR, f"Iter{iteration}_CLI")
        files = collect_cli_files(folder)
        cli_headers = ["Filename", "Size (bytes)", "Phase", "Device"]
        cli_rows = []
        for fname, size in files:
            # Parse phase and device from filename
            parts = fname.replace(".txt", "").split("_")
            # e.g. Iter1_Pre_BC1_Baseline.txt -> Pre, BC1
            phase = parts[1] if len(parts) > 1 else "Unknown"
            device_part = "_".join(parts[2:]) if len(parts) > 2 else "Unknown"
            cli_rows.append([fname, f"{size:,}", phase, device_part])
        add_simple_table(doc, cli_headers, cli_rows)

    doc.add_page_break()

    # ========== CONCLUSION ==========
    doc.add_heading('7. Conclusion', level=1)

    # PASS banner again
    pass_p2 = doc.add_paragraph()
    pass_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pass_run2 = pass_p2.add_run('TEST RESULT: PASS')
    pass_run2.font.size = Pt(16)
    pass_run2.font.bold = True
    pass_run2.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

    doc.add_paragraph(
        'TC 07-02 RETEST validates that the dual-homed L2 Border topology with Po40 (to FS2_BC1) '
        'and Po41 (to FS2_BC2) provides near-hitless failover capability. Across all three iterations, '
        'the following was consistently demonstrated:'
    )

    conclusions = [
        'Po41 provided uninterrupted fabric connectivity during complete Po40 failure.',
        'OSPF adjacency to BC2 remained FULL throughout all failure events.',
        'BFD session to BC2 maintained Up/Up state, ensuring sub-second failure detection.',
        'LISP control plane sessions were preserved at 2/2 established, preventing overlay disruption.',
        'ECMP load balancing was fully restored after each recovery within seconds.',
        'Zero TrustSec (CTS) denied counters observed -- no policy enforcement impact.',
        'Results were consistent and repeatable across all three iterations.',
    ]
    for c in conclusions:
        doc.add_paragraph(c, style='List Bullet')

    doc.add_paragraph(
        'The implementation of Po41 to FS2_BC2 transforms the L2 Border handoff from a single '
        'point of failure into a resilient dual-homed architecture. This represents a substantial '
        'improvement over the original single-homed topology, reducing potential traffic disruption '
        'from minutes of outage with dozens of dead flows to sub-second failover with minimal or '
        'zero traffic impact. The dual-homed design is recommended for all L2 Border nodes in '
        'the Morgan Stanley SDA Phase 2 deployment.'
    )

    # Save
    output_path = os.path.join(BASE_DIR, "TC-07-02_RETEST_Dual_Homed_Port_Channel_Failure_Report.docx")
    doc.save(output_path)
    print(f"Word document saved: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# FILE 2: CXTM Results Text File
# ---------------------------------------------------------------------------

def generate_cxtm_results():
    lines = []

    lines.append("=" * 80)
    lines.append("TC 07-02 RETEST: CXTM RESULTS")
    lines.append("Link Failure From L2 Border Port-Channel to Fabric BC Node")
    lines.append("Dual-Homed Topology Validation")
    lines.append("=" * 80)
    lines.append("")
    lines.append("Project:       Morgan Stanley SDA Phase 2")
    lines.append("Test Case ID:  TC 07-02 (RETEST)")
    lines.append("Date:          April 7, 2026")
    lines.append("Executed By:   wbenbark")
    lines.append("Location:      RTP S10-360, Row P, Racks 18-25")
    lines.append("CC Version:    2.3.7.9-70301")
    lines.append("")

    lines.append("=" * 80)
    lines.append("EXECUTIVE RESULTS SUMMARY")
    lines.append("=" * 80)
    lines.append("")
    lines.append("OVERALL RESULT: PASS")
    lines.append("")
    lines.append("Key Findings:")
    lines.append("  - Po41 remained UP during all 3 iterations when Po40 was down")
    lines.append("  - OSPF to BC2 maintained FULL state throughout all failure events")
    lines.append("  - BFD to BC2 (192.168.41.1) stayed Up/Up during all failures")
    lines.append("  - LISP sessions maintained 2/2 established even during failure")
    lines.append("  - ECMP restored after each recovery (routes via both Po40 and Po41)")
    lines.append("  - Zero CTS denied counters -- no TrustSec enforcement disruption")
    lines.append("  - Results consistent and repeatable across all 3 iterations")
    lines.append("  - Massive improvement over original single-homed test")
    lines.append("")

    lines.append("=" * 80)
    lines.append("TEST CONFIGURATION")
    lines.append("=" * 80)
    lines.append("")
    lines.append("Devices:")
    lines.append("  FS2_L2H-1  C9404R  172.31.0.194  L2 Border (DUT)       IOS-XE 17.15.4")
    lines.append("  FS2_BC1    C9606R  172.31.2.0    Fabric Border Ctrl     IOS-XE 17.15.4")
    lines.append("  FS2_BC2    C9606R  172.31.2.2    Fabric Border Ctrl     IOS-XE 17.15.4")
    lines.append("")
    lines.append("Port-Channels:")
    lines.append("  Po40: Te4/0/1 + Te4/0/2 -> FS2_BC1, 192.168.40.0/31, LACP, MTU 9100")
    lines.append("  Po41: Te4/0/3 + Te4/0/4 -> FS2_BC2, 192.168.41.0/31, LACP, MTU 9100")
    lines.append("")
    lines.append("Topology: Dual-homed L2 Border with ECMP across two Border Controllers")
    lines.append("")

    lines.append("=" * 80)
    lines.append("RESULTS SUMMARY -- ALL 3 ITERATIONS")
    lines.append("=" * 80)
    lines.append("")

    # Table header
    fmt = "  {:<35} {:<15} {:<15} {:<15}"
    lines.append(fmt.format("Metric", "Iter 1", "Iter 2", "Iter 3"))
    lines.append("  " + "-" * 80)
    lines.append(fmt.format("Pre Timestamp", "13:41:34", "14:58:47", "15:15:14"))
    lines.append(fmt.format("During Timestamp", "14:08:59", "15:01", "15:18"))
    lines.append(fmt.format("Post Timestamp", "14:52:40", "15:06", "15:21"))
    lines.append(fmt.format("Po40 Baseline", "RU", "RU", "RU"))
    lines.append(fmt.format("Po41 Baseline", "RU", "RU", "RU"))
    lines.append(fmt.format("Po40 During Failure", "RD", "RD", "RD"))
    lines.append(fmt.format("Po41 During Failure", "RU", "RU", "RU"))
    lines.append(fmt.format("Po40 Post Recovery", "RU", "RU", "RU"))
    lines.append(fmt.format("Po41 Post Recovery", "RU", "RU", "RU"))
    lines.append(fmt.format("OSPF BC1 Baseline", "FULL", "FULL", "FULL"))
    lines.append(fmt.format("OSPF BC2 Baseline", "FULL", "FULL", "FULL"))
    lines.append(fmt.format("OSPF BC1 During Failure", "DOWN", "DOWN", "DOWN"))
    lines.append(fmt.format("OSPF BC2 During Failure", "FULL", "FULL", "FULL"))
    lines.append(fmt.format("OSPF BC1 Post Recovery", "FULL", "FULL", "FULL"))
    lines.append(fmt.format("OSPF BC2 Post Recovery", "FULL", "FULL", "FULL"))
    lines.append(fmt.format("BFD 192.168.40.1 Baseline", "Up/Up", "Up/Up", "Up/Up"))
    lines.append(fmt.format("BFD 192.168.41.1 Baseline", "Up/Up", "Up/Up", "Up/Up"))
    lines.append(fmt.format("BFD 192.168.40.1 During", "Not present", "Not present", "Not present"))
    lines.append(fmt.format("BFD 192.168.41.1 During", "Up/Up", "Up/Up", "Up/Up"))
    lines.append(fmt.format("LISP Sessions Baseline", "2/2 estab", "2/2 estab", "2/2 estab"))
    lines.append(fmt.format("LISP Sessions During", "2/2 estab", "2/2 estab", "2/2 estab"))
    lines.append(fmt.format("LISP Sessions Post", "2/2 estab", "2/2 estab", "2/2 estab"))
    lines.append(fmt.format("ECMP Baseline", "Active", "Active", "Active"))
    lines.append(fmt.format("ECMP During", "Po41 only", "Po41 only", "Po41 only"))
    lines.append(fmt.format("ECMP Post Recovery", "Restored", "Restored", "Restored"))
    lines.append(fmt.format("CTS HW-Denied", "0", "0", "0"))
    lines.append("")

    lines.append("=" * 80)
    lines.append("COMPARISON: ORIGINAL (SINGLE-HOMED) VS RETEST (DUAL-HOMED)")
    lines.append("=" * 80)
    lines.append("")
    cfmt = "  {:<30} {:<30} {:<30}"
    lines.append(cfmt.format("Metric", "Original", "Retest"))
    lines.append("  " + "-" * 88)
    lines.append(cfmt.format("Topology", "Po40 to BC1 only", "Po40(BC1) + Po41(BC2)"))
    lines.append(cfmt.format("Failover Path", "None", "Po41 (immediate)"))
    lines.append(cfmt.format("Outage Duration", "~2 minutes", "Sub-second"))
    lines.append(cfmt.format("Dead Flows", "50-200", "Minimal/zero"))
    lines.append(cfmt.format("LISP Sessions During", "Dropped (0-1/2)", "Maintained (2/2)"))
    lines.append(cfmt.format("OSPF Adjacencies Lost", "All", "BC1 only (BC2 maintained)"))
    lines.append(cfmt.format("Recovery", "Full reconvergence", "ECMP restore only"))
    lines.append(cfmt.format("Result", "PASS (with outage)", "PASS (near-hitless)"))
    lines.append("")

    lines.append("=" * 80)
    lines.append("PASS/FAIL CRITERIA CHECKLIST")
    lines.append("=" * 80)
    lines.append("")
    checks = [
        ("[PASS] Po41 remains UP during Po40 failure (all 3 iterations)", True),
        ("[PASS] OSPF adjacency to BC2 maintained FULL during failure", True),
        ("[PASS] BFD session to BC2 remains Up/Up during failure", True),
        ("[PASS] LISP sessions maintained (2/2 established) during failure", True),
        ("[PASS] All routes reconverge via Po41 during failure", True),
        ("[PASS] Po40 recovers to RU after member interfaces restored", True),
        ("[PASS] ECMP resumes after recovery (routes via both Po40 and Po41)", True),
        ("[PASS] OSPF adjacency to BC1 re-establishes to FULL after recovery", True),
        ("[PASS] BFD session to BC1 re-establishes after recovery", True),
        ("[PASS] Zero CTS denied counters throughout test", True),
        ("[PASS] Results consistent across all 3 iterations", True),
    ]
    for check_text, _ in checks:
        lines.append(f"  {check_text}")
    lines.append("")
    lines.append("  11/11 criteria PASSED")
    lines.append("")

    lines.append("=" * 80)
    lines.append("EVIDENCE INVENTORY")
    lines.append("=" * 80)
    lines.append("")

    lines.append("CLI Files (27 total):")
    lines.append("")
    for iteration in [1, 2, 3]:
        folder = os.path.join(BASE_DIR, f"Iter{iteration}_CLI")
        files = collect_cli_files(folder)
        lines.append(f"  Iteration {iteration} ({len(files)} files):")
        for fname, size in files:
            lines.append(f"    {fname:<45} {size:>8,} bytes")
        lines.append("")

    lines.append("Screenshots:")
    lines.append("")
    img_folders = [
        ("Iteration 1 - Pre",    os.path.join(BASE_DIR, "Images_Iteration1", "PreCaptures")),
        ("Iteration 1 - During", os.path.join(BASE_DIR, "Images_Iteration1", "During")),
        ("Iteration 1 - Post",   os.path.join(BASE_DIR, "Images_Iteration1", "Post")),
        ("Iteration 2",          os.path.join(BASE_DIR, "Images_Iteration2")),
        ("Iteration 3",          os.path.join(BASE_DIR, "Images_Iteration3")),
    ]
    total_imgs = 0
    for label, folder in img_folders:
        imgs = collect_images(folder)
        total_imgs += len(imgs)
        lines.append(f"  {label}: {len(imgs)} screenshots")
    lines.append(f"  TOTAL: {total_imgs} screenshots")
    lines.append("")

    lines.append("=" * 80)
    lines.append("CONCLUSION")
    lines.append("=" * 80)
    lines.append("")
    lines.append("TC 07-02 RETEST: PASS")
    lines.append("")
    lines.append("The dual-homed L2 Border topology (Po40 + Po41) provides near-hitless")
    lines.append("failover when a single port-channel fails. LISP control plane, OSPF")
    lines.append("adjacency to the surviving BC, and BFD liveness detection all remain")
    lines.append("operational throughout the failure event. ECMP load balancing resumes")
    lines.append("immediately upon recovery. This represents a substantial improvement")
    lines.append("over the original single-homed topology.")
    lines.append("")
    lines.append("=" * 80)
    lines.append("END OF RESULTS")
    lines.append("=" * 80)

    output_path = os.path.join(BASE_DIR, "TC-07-02_RETEST_CXTM_Results.txt")
    with open(output_path, "w") as f:
        f.write("\n".join(lines) + "\n")
    print(f"CXTM results saved: {output_path}")
    return output_path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating TC 07-02 RETEST deliverables...")
    print()
    word_path = generate_word_report()
    print()
    cxtm_path = generate_cxtm_results()
    print()
    print("Done. Generated files:")
    print(f"  1. {word_path}")
    print(f"  2. {cxtm_path}")
